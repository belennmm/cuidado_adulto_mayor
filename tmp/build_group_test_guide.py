from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"C:\Users\leejo\Desktop\cuidado_adulto_mayor\output\documentos\Guion_Grupal_Ejecucion_Pruebas.docx"
NAVY, BLUE, GRAY, WHITE = "18324A", "2E74B5", "5B6573", "FFFFFF"
LIGHT, PALE, GREEN, BLACK = "F2F4F7", "EAF2F8", "E7F4EC", "111111"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.49)

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=BLACK):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name); rpr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

normal = doc.styles["Normal"]
normal.font.name, normal.font.size = "Calibri", Pt(11)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, NAVY, 8, 4)]:
    s=doc.styles[name]; s.font.name="Calibri"; s.font.size=Pt(size); s.font.bold=True
    s.font.color.rgb=RGBColor.from_string(color)
    s.paragraph_format.space_before=Pt(before); s.paragraph_format.space_after=Pt(after)
    s.paragraph_format.keep_with_next=True

def shade(cell, fill):
    pr=cell._tc.get_or_add_tcPr(); sh=pr.find(qn("w:shd"))
    if sh is None: sh=OxmlElement("w:shd"); pr.append(sh)
    sh.set(qn("w:fill"),fill)

def border(cell, color="CBD5E1"):
    pr=cell._tc.get_or_add_tcPr(); bs=pr.first_child_found_in("w:tcBorders")
    if bs is None: bs=OxmlElement("w:tcBorders"); pr.append(bs)
    for edge in ("top","left","bottom","right"):
        el=bs.find(qn("w:"+edge))
        if el is None: el=OxmlElement("w:"+edge); bs.append(el)
        el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"6"); el.set(qn("w:color"),color)

def box(title, body, fill=PALE):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shade(c,fill); border(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(2)
    set_font(p.add_run(title+"\n"),size=10.5,bold=True,color=NAVY)
    set_font(p.add_run(body),size=10)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def command(text):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shade(c,"F7F8FA"); border(c,"D9DEE5")
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    set_font(p.add_run(text),"Consolas",9,color="263238")
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def bullet(text):
    p=doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent=Inches(.5); p.paragraph_format.first_line_indent=Inches(-.25)
    p.paragraph_format.space_after=Pt(4); p.add_run(text)

def quote(text):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.28)
    p.paragraph_format.right_indent=Inches(.18); p.paragraph_format.space_after=Pt(8)
    r=p.add_run('“'+text+'”'); set_font(r,size=10.5,italic=True,color="374151")

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("CC3091  |  Demostración de pruebas automatizadas"),size=8.5,color=GRAY)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
set_font(footer.add_run("Página "),size=9,color=GRAY)
field=OxmlElement("w:fldSimple"); field.set(qn("w:instr"),"PAGE"); footer._p.append(field)

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(90); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("GUION GRUPAL"),size=12,bold=True,color=BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("EJECUCIÓN DE PRUEBAS\nAUTOMATIZADAS"),size=27,bold=True,color=NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("Sistema de cuidado del adulto mayor"),size=15,color=GRAY)
box("Distribución", "Chon, Wichandro, Belén y Sebas\nDos grupos de pruebas por integrante: uno de frontend y uno de backend.")
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(70)
set_font(p.add_run("CC3091 - Ingeniería de Software 2\nJulio de 2026"),size=11,color=GRAY)
doc.add_page_break()

people = [
 ("CHON", "Gestión de sesión + inventario de medicamentos", [
  ("1. Gestión de sesión por rol", "Frontend - Vitest", r"cd C:\Users\leejo\Desktop\cuidado_adulto_mayor\frontend\nnpm.cmd test -- tests/unit/auth-session.test.js",
   "Comprueba que el frontend guarde y recupere el token y los datos del usuario según su rol, y que pueda eliminar la sesión.",
   "Mi primera prueba verifica la gestión de sesiones en el frontend. Se comprueba que el sistema guarde el token y los datos del usuario, respete el rol asignado y pueda eliminar la sesión.",
   "3 pruebas aprobadas"),
  ("2. Estado del inventario de medicamentos", "Backend - PHPUnit", r"cd C:\Users\leejo\Desktop\cuidado_adulto_mayor\backend\nphp artisan test --filter=MedicationInventoryStatusTest",
   "Valida los estados vencido, próximo a vencer, inventario bajo y disponible mediante cuatro conjuntos de datos.",
   "Mi segunda prueba verifica las reglas del inventario de medicamentos. Se prueban los estados vencido, próximo a vencer, inventario bajo y disponible.",
   "4 pruebas aprobadas")]),
 ("WICHANDRO", "Configuración de API + autenticación", [
  ("1. Construcción de URL y encabezados", "Frontend - Vitest", r"cd C:\Users\leejo\Desktop\cuidado_adulto_mayor\frontend\nnpm.cmd test -- tests/unit/api-config.test.js",
   "Comprueba URLs relativas y absolutas, encabezados JSON, token Bearer, respuestas JSON y errores HTTP.",
   "Mi primera prueba corresponde a la configuración de la API. Se verifica la construcción de URLs, los encabezados JSON, el token de autenticación y el manejo de respuestas y errores.",
   "4 pruebas aprobadas"),
  ("2. Autenticación y tokens", "Backend - PHPUnit", r"cd C:\Users\leejo\Desktop\cuidado_adulto_mayor\backend\nphp artisan test --filter=AuthenticationTest",
   "Verifica login, registro, credenciales incorrectas, cuenta pendiente, perfil y revocación del token al cerrar sesión.",
   "Mi segunda prueba verifica la autenticación del backend. Se comprueba el inicio de sesión, la generación de tokens con Sanctum, el registro, el acceso al perfil y la revocación del token.",
   "8 pruebas aprobadas")]),
 ("BELÉN", "Inicio de sesión + rutinas", [
  ("1. Inicio de sesión", "Frontend - Vitest", r"cd C:\Users\leejo\Desktop\cuidado_adulto_mayor\frontend\nnpm.cmd test -- tests/dom/login.test.js",
   "Simula el formulario y valida campos faltantes, visibilidad de contraseña, sesión, redirección y errores de API.",
   "Mi primera prueba simula el formulario de inicio de sesión. Se verifica la validación de credenciales, el envío a la API, el almacenamiento de la sesión y la redirección según el rol.",
   "4 pruebas aprobadas"),
  ("2. Rutinas del adulto mayor", "Backend - PHPUnit", r"cd C:\Users\leejo\Desktop\cuidado_adulto_mayor\backend\nphp artisan test --filter=RutinaEndpointsTest",
   "Cubre creación, consulta, modificación, eliminación, actividades completadas, horarios y permisos de acceso.",
   "Mi segunda prueba verifica el módulo de rutinas. Se prueban las operaciones CRUD, las actividades completadas, las validaciones de horario y las restricciones por asignación.",
   "14 pruebas aprobadas")]),
 ("SEBAS", "Registro + administración de usuarios", [
  ("1. Registro de usuarios", "Frontend - Vitest", r"cd C:\Users\leejo\Desktop\cuidado_adulto_mayor\frontend\nnpm.cmd test -- tests/dom/register.test.js",
   "Comprueba campos obligatorios, payload enviado, respuesta exitosa, visualización de contraseña y errores del backend.",
   "Mi primera prueba corresponde al formulario de registro. Se comprueba la validación de campos, el envío de datos y los mensajes de éxito o error.",
   "3 pruebas aprobadas"),
  ("2. Administración de usuarios", "Backend - PHPUnit", r"cd C:\Users\leejo\Desktop\cuidado_adulto_mayor\backend\nphp artisan test --filter=AdminUserManagementTest",
   "Verifica permisos y operaciones para crear, consultar, actualizar, aprobar, rechazar y eliminar usuarios.",
   "Mi segunda prueba verifica la administración de usuarios. Se comprueba el acceso exclusivo del administrador, las operaciones de gestión y las validaciones de seguridad.",
   "5 pruebas aprobadas")])]

for pi,(name,subtitle,tests) in enumerate(people):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(3)
    set_font(p.add_run(name),size=22,bold=True,color=NAVY)
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(10)
    set_font(p.add_run(subtitle),size=12.5,color=GRAY)
    for title,layer,cmd,desc,speech,result in tests:
        doc.add_heading(title,level=2)
        p=doc.add_paragraph(); set_font(p.add_run(layer),size=9.5,bold=True,color=BLUE)
        doc.add_paragraph(desc)
        command(cmd.replace(r"\n","\n"))
        p=doc.add_paragraph(); set_font(p.add_run("Guion para decir:"),size=10,bold=True,color=NAVY)
        quote(speech)
        box("Resultado esperado", result, GREEN)
    if pi < len(people)-1: doc.add_page_break()

doc.add_page_break()
doc.add_heading("Cierre grupal",level=1)
doc.add_paragraph("Al terminar las ocho demostraciones, una persona del grupo puede realizar el cierre.")
quote("En total demostramos ocho grupos de pruebas automatizadas: cuatro del frontend con Vitest y cuatro del backend con PHPUnit. Estas pruebas verifican sesiones, API, formularios, autenticación, inventario, rutinas y administración de usuarios. En la ejecución completa se aprobaron 14 pruebas de frontend y 88 de backend, para un total de 102 pruebas sin fallos.")
doc.add_heading("Orden recomendado para la grabación",level=2)
for x in ["Chon: sesión e inventario.","Wichandro: API y autenticación.","Belén: inicio de sesión y rutinas.","Sebas: registro y administración.","Cierre grupal y presentación del resultado total."]: bullet(x)
box("Recomendación", "Antes de grabar, cada integrante debe ejecutar sus dos comandos una vez. Durante el video, mostrar primero el archivo de prueba, explicar brevemente su objetivo y luego ejecutar el comando en PowerShell.")
doc.add_heading("Resultados globales esperados",level=2)
box("Frontend", "4 archivos, 14 pruebas aprobadas.", GREEN)
box("Backend", "88 pruebas y 405 aserciones aprobadas.", GREEN)

doc.core_properties.title="Guion grupal para ejecución de pruebas automatizadas"
doc.core_properties.author="Equipo de desarrollo"
doc.core_properties.subject="Vitest y PHPUnit - Sistema de cuidado del adulto mayor"
doc.save(OUT)
print(OUT)
