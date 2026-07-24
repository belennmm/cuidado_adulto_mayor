from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK

OUT = r"C:\Users\leejo\Desktop\cuidado_adulto_mayor\output\documentos\Informe_Pruebas_Automatizadas.docx"
NAVY = "18324A"
BLUE = "2E74B5"
SKY = "EAF2F8"
LIGHT = "F2F4F7"
GREEN = "1F6B4A"
GRAY = "5B6573"
WHITE = "FFFFFF"
BLACK = "111111"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.49)

def font(run, name="Calibri", size=11, bold=False, italic=False, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold, run.italic = bold, italic
    run.font.color.rgb = RGBColor.from_string(color)

styles = doc.styles
normal = styles["Normal"]
normal.font.name, normal.font.size = "Calibri", Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for style_name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, NAVY, 8, 4)]:
    s = styles[style_name]
    s.font.name, s.font.size, s.font.bold = "Calibri", Pt(size), True
    s.font.color.rgb = RGBColor.from_string(color)
    s.paragraph_format.space_before, s.paragraph_format.space_after = Pt(before), Pt(after)
    s.paragraph_format.keep_with_next = True

code_style = styles.add_style("Codigo", WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name, code_style.font.size = "Consolas", Pt(8.2)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(8)
code_style.paragraph_format.left_indent = Inches(0.18)
code_style.paragraph_format.right_indent = Inches(0.18)
code_style.paragraph_format.line_spacing = 1.0

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None: tcPr.append(shd)

def borders(cell, color="D4DAE2", size="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders"); tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tcBorders.find(qn(f"w:{edge}")) or OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), size); el.set(qn("w:color"), color)
        if el.getparent() is None: tcBorders.append(el)

def set_cell_text(cell, text, bold=False, color=BLACK, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    for i, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[i], h, True, WHITE, 9)
        shade(t.rows[0].cells[i], NAVY); borders(t.rows[0].cells[i], NAVY)
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), False, BLACK, 9)
            shade(cells[i], WHITE if ri % 2 == 0 else LIGHT); borders(cells[i])
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths): row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(.5); p.paragraph_format.first_line_indent = Inches(-.25)
    p.paragraph_format.space_after = Pt(5)
    p.add_run(text)

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(.5); p.paragraph_format.first_line_indent = Inches(-.25)
    p.paragraph_format.space_after = Pt(5)
    p.add_run(text)

def callout(title, body, color=SKY):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0,0); shade(c, color); borders(c, "B8C9DA")
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(3)
    font(p.add_run(title + "\n"), size=10.5, bold=True, color=NAVY)
    font(p.add_run(body), size=10, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def code(text):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shade(c,"F7F8FA"); borders(c,"D9DEE5")
    p=c.paragraphs[0]; p.style=code_style
    r=p.add_run(text.strip()); font(r,"Consolas",8.2,color="263238")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página "); font(run,size=9,color=GRAY)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)

# Running furniture
header = sec.header.paragraphs[0]
header.text = "CC3091 - Ingeniería de Software 2  |  Pruebas automatizadas"
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for r in header.runs: font(r,size=8.5,color=GRAY)
add_page_number(sec.footer.paragraphs[0])

# Cover
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(80); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("UNIVERSIDAD DEL VALLE DE GUATEMALA"),size=11,bold=True,color=BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(18)
font(p.add_run("PRUEBAS\nAUTOMATIZADAS"),size=29,bold=True,color=NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
font(p.add_run("Informe técnico de implementación y resultados"),size=15,color=GRAY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(30)
font(p.add_run("Sistema de cuidado del adulto mayor"),size=14,bold=True,color=BLUE)
table(["Curso", "Entrega", "Equipo"], [["CC3091 - Ingeniería de Software 2", "Tarea 3 - Semestre II 2026", "[Agregar nombres y carnés]"]], [2.3,2.0,2.2])
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(65)
font(p.add_run("Guatemala, julio de 2026"),size=11,color=GRAY)
doc.add_page_break()

doc.add_heading("Resumen ejecutivo", level=1)
doc.add_paragraph("Se investigaron alternativas para automatizar pruebas en el frontend JavaScript y el backend PHP/Laravel del sistema de cuidado del adulto mayor. Se seleccionaron Vitest con jsdom y Testing Library para el frontend, y PHPUnit integrado con Laravel para el backend. La implementación verificada comprende 102 pruebas: 14 de frontend y 88 de backend, todas aprobadas.")
table(["Capa", "Herramientas", "Archivos", "Pruebas", "Resultado", "Tiempo observado"], [
    ["Frontend", "Vitest 3.2.7 + jsdom + Testing Library", "4", "14", "14 aprobadas", "1.45 s"],
    ["Backend", "PHPUnit 11.5 + Laravel Test", "17", "88 (405 aserciones)", "88 aprobadas", "1.67 s"],
    ["Total", "Dos suites automatizadas", "21", "102", "0 fallos", "3.12 s"]], [1.0,1.75,.65,.7,1.05,1.1])
callout("Resultado principal", "Las dos suites finalizaron sin errores. Los tiempos corresponden a las ejecuciones realizadas el 19 de julio de 2026 en el equipo local y pueden variar según el entorno.")

doc.add_heading("1. Objetivo y alcance", level=1)
doc.add_paragraph("El objetivo fue seleccionar herramientas apropiadas para la tecnología del proyecto, implementar una muestra representativa de pruebas automatizadas y evaluar su utilidad. El alcance incluye lógica unitaria, manipulación del DOM, comunicación con la API, autenticación, autorización, validaciones, persistencia y reglas de negocio.")
doc.add_heading("Competencias cubiertas", level=2)
bullet("Investigación de frameworks para JavaScript y PHP/Laravel.")
bullet("Selección razonada de Vitest y PHPUnit según compatibilidad, rapidez, documentación y robustez.")
bullet("Implementación de más de tres ejemplos en frontend y backend.")
bullet("Análisis de resultados, ventajas, desventajas y experiencia de uso.")

doc.add_heading("2. Herramientas investigadas", level=1)
doc.add_heading("2.1 Frontend JavaScript", level=2)
table(["Herramienta", "Fortalezas", "Limitaciones", "Adecuación"], [
    ["Vitest", "ESM, mocks, modo watch, API familiar y ejecución paralela.", "El DOM requiere jsdom o navegador; ecosistema más joven que Jest.", "Muy alta"],
    ["Jest", "Ecosistema maduro, amplia documentación, mocks y snapshots.", "Puede requerir más configuración para ESM y flujos modernos.", "Alta"],
    ["Mocha + Chai", "Flexible y modular.", "Exige escoger y configurar varias piezas por separado.", "Media"],
    ["Playwright/Cypress", "Prueban flujos en navegadores reales.", "Mayor costo y tiempo; su foco principal es E2E, no la unidad.", "Complementaria"]], [1.1,2.15,2.15,1.1])
doc.add_paragraph("Testing Library se utiliza junto con Vitest para interactuar con formularios como lo haría una persona usuaria. jsdom aporta un DOM simulado y permite probar la interfaz rápidamente sin abrir un navegador.")

doc.add_heading("2.2 Backend PHP/Laravel", level=2)
table(["Herramienta", "Fortalezas", "Limitaciones", "Adecuación"], [
    ["PHPUnit", "Estándar de facto en PHP; aserciones, data providers, mocks y cobertura.", "Puede ser verboso en suites extensas.", "Muy alta"],
    ["Pest", "Sintaxis concisa y compatible con PHPUnit.", "Añade una capa y convenciones adicionales.", "Alta"],
    ["Codeception", "Une pruebas unitarias, funcionales y de aceptación.", "Configuración y alcance mayores de lo necesario.", "Media"],
    ["Laravel Dusk", "Automatización de navegador para aplicaciones Laravel.", "Más lento; adecuado para E2E, no para lógica aislada.", "Complementaria"]], [1.1,2.15,2.15,1.1])

doc.add_heading("3. Herramientas seleccionadas y justificación", level=1)
doc.add_heading("3.1 Vitest + jsdom + Testing Library", level=2)
bullet("Compatibilidad directa con módulos ES usados por el frontend y configuración sencilla mediante vitest.config.js.")
bullet("Mocks integrados con vi.fn() y vi.stubGlobal(), útiles para aislar fetch, navegación y almacenamiento de sesión.")
bullet("Ejecución rápida y modo watch, que vuelve a lanzar las pruebas al modificar código.")
bullet("API compatible con el estilo de Jest, lo que reduce la curva de aprendizaje.")
bullet("Testing Library favorece pruebas centradas en comportamiento observable y no en detalles internos.")
doc.add_heading("3.2 PHPUnit integrado con Laravel", level=2)
bullet("Ya forma parte del ecosistema Laravel y se ejecuta con php artisan test.")
bullet("Incluye clientes de prueba HTTP, aserciones JSON y aserciones sobre la base de datos.")
bullet("RefreshDatabase y SQLite en memoria aíslan cada caso y evitan contaminar datos reales.")
bullet("Sanctum::actingAs permite comprobar autenticación y roles con bajo costo de configuración.")
bullet("Los data providers reducen duplicación cuando una regla tiene varias entradas y salidas.")
callout("Decisión", "Vitest y PHPUnit ofrecen el mejor equilibrio para la arquitectura actual: ejecución rápida, integración natural con el código existente, buena documentación y suficiente capacidad para crecer hacia cobertura y CI.")

doc.add_heading("4. Implementación en el proyecto", level=1)
doc.add_heading("4.1 Configuración del frontend", level=2)
code('''// frontend/vitest.config.js\nexport default defineConfig({\n  test: { environment: "jsdom", setupFiles: ["./tests/setup.js"],\n          clearMocks: true, restoreMocks: true }\n})''')
doc.add_paragraph("El comando npm.cmd test ejecuta vitest run una sola vez. npm.cmd run test:watch mantiene la suite activa durante el desarrollo.")
doc.add_heading("4.2 Configuración del backend", level=2)
code('''<!-- backend/phpunit.xml -->\n<testsuite name="Unit"><directory>tests/Unit</directory></testsuite>\n<testsuite name="Feature"><directory>tests/Feature</directory></testsuite>\n<env name="DB_CONNECTION" value="sqlite"/>\n<env name="DB_DATABASE" value=":memory:"/>''')
doc.add_paragraph("El comando php artisan test ejecuta las suites Unit y Feature usando una base de datos SQLite en memoria.")

doc.add_heading("5. Ejemplos automatizados del frontend", level=1)
examples_front = [
 ("5.1 Gestión de sesión por rol", "Comprueba guardar, recuperar, restringir por rol y eliminar una sesión.", '''window.AuthSession.saveSession("abc123", user)\nexpect(window.AuthSession.getToken(["profesional"])).toBe("abc123")\nexpect(window.AuthSession.getUser(["profesional"])).toEqual(user)''', "tests/unit/auth-session.test.js - 3 pruebas"),
 ("5.2 Construcción de URL y encabezados", "Valida rutas relativas, URLs absolutas y el encabezado Bearer.", '''expect(CuidadoApi.buildUrl("/login"))\n  .toBe("https://api.example.test/api/login")\nexpect(headers.Authorization).toBe("Bearer token-prueba")''', "tests/unit/api-config.test.js - 4 pruebas"),
 ("5.3 Inicio de sesión", "Simula el formulario, aísla la API y confirma mensaje, sesión y redirección según rol.", '''fetchJson.mockResolvedValue({ token: "token", user: { role: "admin" } })\nfireEvent.click(getByRole(document.body, "button", { name: "Ingresar" }))\nexpect(saveSession).toHaveBeenCalled()''', "tests/dom/login.test.js - 4 pruebas"),
 ("5.4 Registro de usuario", "Comprueba campos obligatorios, payload enviado, respuesta exitosa y error del backend.", '''fireEvent.submit(document.querySelector("form"))\nexpect(fetchJson).toHaveBeenCalledOnce()\nexpect(JSON.parse(options.body)).toMatchObject({ role: "familiar" })''', "tests/dom/register.test.js - 3 pruebas")]
for title, purpose, snippet, evidence in examples_front:
    doc.add_heading(title, level=2); doc.add_paragraph(purpose); code(snippet)
    p=doc.add_paragraph(); font(p.add_run("Evidencia: "),bold=True,color=GREEN); p.add_run(evidence)

doc.add_heading("6. Ejemplos automatizados del backend", level=1)
examples_back = [
 ("6.1 Estado del inventario de medicamentos", "Un data provider cubre medicamento vencido, próximo a vencer, inventario bajo y disponible.", '''#[DataProvider('inventoryCases')]\n$this->assertSame($expectedKey,\n  $medication->inventoryStatus($today)['key']);''', "MedicationInventoryStatusTest.php - 4 conjuntos de datos"),
 ("6.2 Autenticación y tokens", "Verifica login exitoso, credenciales incorrectas, cuenta pendiente, perfil y revocación de tokens.", '''$this->postJson('/api/login', $credentials)\n  ->assertOk()->assertJsonStructure(['token', 'user']);\n$this->assertDatabaseCount('personal_access_tokens', 1);''', "AuthenticationTest.php - 8 pruebas"),
 ("6.3 Rutinas del adulto mayor", "Cubre CRUD, horarios inválidos, actividades completadas y restricciones por asignación.", '''$this->postJson('/api/rutinas', $payload)\n  ->assertCreated()->assertJsonPath('rutina.horario', '08:00');\n$this->assertDatabaseHas('rutinas', ['nombre' => 'Rutina matutina']);''', "RutinaEndpointsTest.php - 14 pruebas"),
 ("6.4 Administración de usuarios", "Valida autorización, creación, edición, aprobación, eliminación y rechazo de solicitudes.", '''Sanctum::actingAs($admin);\n$this->patchJson("/api/admin/users/{$userId}/approve")\n  ->assertOk()->assertJsonPath('user.is_approved', true);''', "AdminUserManagementTest.php - 5 pruebas")]
for title, purpose, snippet, evidence in examples_back:
    doc.add_heading(title, level=2); doc.add_paragraph(purpose); code(snippet)
    p=doc.add_paragraph(); font(p.add_run("Evidencia: "),bold=True,color=GREEN); p.add_run(evidence)

doc.add_heading("7. Resultados de ejecución", level=1)
table(["Suite", "Comando", "Archivos", "Pruebas/aserciones", "Estado", "Duración"], [
 ["Frontend", "npm.cmd test", "4", "14 pruebas", "PASS", "1.45 s"],
 ["Backend", "php artisan test", "17", "88 pruebas / 405 aserciones", "PASS", "1.67 s"]], [1.0,1.35,.7,1.55,.75,1.0])
code('''FRONTEND\nTest Files  4 passed (4)\nTests       14 passed (14)\nDuration    1.45s\n\nBACKEND\nTests:      88 passed (405 assertions)\nDuration:   1.67s''')
doc.add_paragraph("La evidencia demuestra que las dependencias están instaladas, las configuraciones son válidas y todos los escenarios implementados se ejecutan de forma repetible. No se observaron fallos en las dos corridas reportadas.")

doc.add_heading("8. Ventajas, desventajas y experiencia", level=1)
table(["Aspecto", "Vitest", "PHPUnit/Laravel"], [
 ["Ventajas", "Rápido, mocks integrados, DOM simulado, modo watch.", "HTTP/JSON/BD integrados, aislamiento, soporte nativo de Laravel."],
 ["Desventajas", "jsdom no reemplaza por completo un navegador real.", "Las pruebas Feature pueden crecer en costo y tiempo."],
 ["Escritura", "beforeEach y mocks facilitan escenarios de interfaz.", "Factories y RefreshDatabase simplifican estados reproducibles."],
 ["Mantenimiento", "Selectores accesibles y contratos públicos reducen fragilidad.", "Conviene reutilizar factories y evitar fixtures excesivos."]], [1.2,2.65,2.65])
doc.add_paragraph("La ejecución fue eficiente: 102 pruebas finalizaron en aproximadamente 3.12 segundos sumando ambas corridas observadas. El mayor esfuerzo no está en ejecutar las herramientas, sino en identificar escenarios, preparar datos y definir aserciones que representen reglas reales del sistema.")

doc.add_heading("9. Conclusiones", level=1)
numbered("La automatización permitió verificar rápidamente reglas críticas de autenticación, autorización, inventario, rutinas y formularios.")
numbered("Vitest se ajusta bien al frontend por su soporte para ES Modules, mocks y DOM simulado; PHPUnit se ajusta al backend por su integración profunda con Laravel.")
numbered("La combinación de pruebas unitarias y Feature ofrece mayor confianza que cualquiera de las dos categorías por separado.")
numbered("Los resultados obtenidos - 102 pruebas aprobadas y cero fallos - establecen una línea base reproducible para detectar regresiones.")
numbered("Como siguiente paso, se recomienda agregar cobertura de código, integrar las suites en CI y complementar los flujos principales con Playwright o Laravel Dusk en un navegador real.")

doc.add_heading("10. Guion sugerido para el video", level=1)
doc.add_paragraph("La entrega solicita un video que evidencie la implementación. El equipo puede grabar una demostración de 4 a 6 minutos con esta secuencia:")
bullet("Mostrar la estructura frontend/tests y backend/tests.")
bullet("Abrir un ejemplo de Vitest y explicar Arrange, Act y Assert.")
bullet("Ejecutar npm.cmd test y mostrar 14 pruebas aprobadas.")
bullet("Abrir un ejemplo de PHPUnit y explicar datos, petición y aserciones.")
bullet("Ejecutar php artisan test y mostrar 88 pruebas y 405 aserciones aprobadas.")
bullet("Cerrar con ventajas, limitaciones y próximos pasos.")
callout("Enlace del video", "[PEGAR AQUÍ EL ENLACE AL VIDEO ANTES DE ENTREGAR EL PDF]", "FFF4CE")

doc.add_heading("Referencias", level=1)
refs = [
 "Vitest. Why Vitest. https://vitest.dev/guide/why.html",
 "Vitest. Getting Started. https://vitest.dev/guide/",
 "Vitest. Features. https://vitest.dev/guide/features.html",
 "PHPUnit. Documentation. https://phpunit.de/documentation.html",
 "PHPUnit. Manual. https://docs.phpunit.de/",
 "Jest. Getting Started. https://jestjs.io/docs/getting-started",
 "Playwright. Test documentation. https://playwright.dev/docs/test-intro",
 "Laravel. HTTP Tests. https://laravel.com/docs/11.x/http-tests",
 "Testing Library. Guiding Principles. https://testing-library.com/docs/guiding-principles"
]
for ref in refs: bullet(ref)

# Keep headings together and add basic document metadata.
doc.core_properties.title = "Pruebas automatizadas - Sistema de cuidado del adulto mayor"
doc.core_properties.subject = "Tarea 3, CC3091 Ingeniería de Software 2"
doc.core_properties.author = "Equipo de desarrollo"
doc.core_properties.keywords = "Vitest, PHPUnit, Laravel, pruebas automatizadas"
doc.save(OUT)
print(OUT)
